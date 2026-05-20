"""
Fix document: image placement, captions, font formatting.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
from lxml import etree

doc = Document(r'C:\Users\KUHN\Desktop\1.docx')

# ============================================
# STEP 1: Fix font on ALL content paragraphs
# ============================================
# Content paragraphs were created by make_paragraph which uses raw XML.
# The eastAsia font was set in rFonts but might not be read correctly.
# Let's also ensure font size is properly set.
fix_count = 0
for p in doc.paragraphs:
    if p.style.name == 'Normal':
        for r in p.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    # Ensure eastAsia is set
                    ea = rFonts.get(qn('w:eastAsia'))
                    if not ea:
                        rFonts.set(qn('w:eastAsia'), '宋体')
                        fix_count += 1
                    # Ensure ascii is set (except code blocks)
                    ascii_f = rFonts.get(qn('w:ascii'))
                    en_f = rFonts.get(qn('w:hAnsi'))
                    if ascii_f == 'Consolas':
                        pass  # Code blocks - keep as is
                    elif not ascii_f or ascii_f == 'Times New Roman':
                        # Already correct, but ensure it's set
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        rFonts.set(qn('w:eastAsia'), '宋体')
                        fix_count += 1

print(f'Font fixes applied: {fix_count}')

# ============================================
# STEP 2: Fix image placement and captions
# ============================================

# Find image paragraphs and their relationships
image_data = []  # (para_index, para_obj, rId, image_name, description_text)

for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//'+qn('w:drawing'))
    for d in drawings:
        blip = d.find('.//'+qn('a:blip'))
        if blip is not None:
            embed = blip.get(qn('r:embed'))
            # Find the image file name from rels
            img_name = ''
            for rel_id, rel in doc.part.rels.items():
                if rel_id == embed and 'image' in rel.reltype:
                    img_name = rel.target_ref
                    break
            image_data.append({
                'idx': i,
                'para': p,
                'rid': embed,
                'img': img_name,
                'drawing': d,
                'para_text': p.text.strip()
            })

print(f'\nImage data:')
for img in image_data:
    print(f'  Para {img["idx"]}: rId={img["rid"]}, img={img["img"]}, text="{img["para_text"][:50]}"')

# Find which description text corresponds to each image
# The descriptions in the text mention 图 5.1 to 图 5.4
# We need to identify which description goes with which image

# Based on the structure, the paragraph ordering should match:
# Para [34] description (1) 图5.1...
# Para [35] IMAGE (empty text) - this is Figure 5.1
# Para [36] description (2) 图5.2... WITH embedded image
# Para [43] description (3) 图5.3... WITH embedded image
# Para [44] description (4) 图5.4... WITH embedded image

# Issue: images 5.2-5.4 are embedded in description paragraphs.
# They should be in their own paragraphs with captions.

# Strategy: 
# 1. For image 5.1 in para [35] (empty para with just image):
#    - The description is in para [34]
#    - I need to make sure there's a caption below the image
#    
# 2. For images 5.2-5.4:
#    - Extract the drawing element from the description paragraph
#    - Insert a new paragraph AFTER the description with just the image
#    - Add a caption below that

# Let me implement this by working with the XML directly

def make_caption_para(text):
    """Create a centered caption paragraph element."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    
    # Center alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    
    # Line spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    
    p.append(pPr)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)
    
    # Bold for captions
    b = OxmlElement('w:b')
    rPr.append(b)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    rPr.append(szCs)
    
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_image_para(rid):
    """Create a paragraph with just an inline image using the existing relationship."""
    # Build the drawing XML for the image by copying from existing
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    
    # Center alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    
    # Line spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    
    p.append(pPr)
    
    # Create the drawing XML
    # Anchor for floating image (type used by Word)
    drawing = OxmlElement('w:drawing')
    
    # wp:inline for inline image
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    
    inline = OxmlElement(qn('wp:inline'))
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')
    
    # extents
    extent = OxmlElement(qn('wp:extent'))
    extent.set('cx', '4000000')  # ~10cm
    extent.set('cy', '3000000')  # ~7.5cm
    inline.append(extent)
    
    # effectExtent
    effExt = OxmlElement(qn('wp:effectExtent'))
    effExt.set('l', '0')
    effExt.set('t', '0')
    effExt.set('r', '0')
    effExt.set('b', '0')
    inline.append(effExt)
    
    # docPr
    docPr = OxmlElement(qn('wp:docPr'))
    docPr.set('id', '1')
    docPr.set('name', 'Picture')
    docPr.set('descr', 'Chart')
    inline.append(docPr)
    
    # graphic
    graphic = OxmlElement(qn('a:graphic'))
    graphic.set('xmlns:a', a_ns)
    
    graphicData = OxmlElement(qn('a:graphicData'))
    graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    
    pic = OxmlElement(qn('pic:pic'))
    pic.set('xmlns:pic', 'http://schemas.openxmlformats.org/drawingml/2006/picture')
    
    # nvPicPr
    nvPicPr = OxmlElement(qn('pic:nvPicPr'))
    cNvPr = OxmlElement(qn('pic:cNvPr'))
    cNvPr.set('id', '0')
    cNvPr.set('name', 'Picture')
    nvPicPr.append(cNvPr)
    cNvPicPr = OxmlElement(qn('pic:cNvPicPr'))
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)
    
    # blipFill
    blipFill = OxmlElement(qn('pic:blipFill'))
    blip = OxmlElement(qn('a:blip'))
    blip.set(qn('r:embed'), rid)
    blip.set('xmlns:r', r_ns)
    blipFill.append(blip)
    stretch = OxmlElement(qn('a:stretch'))
    fillRect = OxmlElement(qn('a:fillRect'))
    stretch.append(fillRect)
    blipFill.append(stretch)
    pic.append(blipFill)
    
    # spPr
    spPr = OxmlElement(qn('pic:spPr'))
    xfrm = OxmlElement(qn('a:xfrm'))
    off = OxmlElement(qn('a:off'))
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext = OxmlElement(qn('a:ext'))
    ext.set('cx', '4000000')
    ext.set('cy', '3000000')
    xfrm.append(ext)
    spPr.append(xfrm)
    prstGeom = OxmlElement(qn('a:prstGeom'))
    prstGeom.set('prst', 'rect')
    spPr.append(prstGeom)
    pic.append(spPr)
    
    graphicData.append(pic)
    graphic.append(graphicData)
    inline.append(graphic)
    drawing.append(inline)
    p.append(drawing)
    
    return p

# Now let's fix images 5.2-5.4
# We need to COPY the drawing element from the mixed paragraph to a new paragraph

# Figure mapping
figures = {
    '5.1': {'desc_idx': 34, 'img_idx': 35},
    '5.2': {'desc_idx': 36, 'img_idx': 36},  # Same as desc - embedded
    '5.3': {'desc_idx': 43, 'img_idx': 43},  # Same as desc - embedded
    '5.4': {'desc_idx': 44, 'img_idx': 44},  # Same as desc - embedded
}

# Actually, let's look at this differently.
# The user inserted images in specific positions.
# For 5.1: image is in its own paragraph [35] - good
# For 5.2-5.4: images are embedded in description text paragraphs

# The best approach: extract the drawing from each mixed paragraph,
# insert a new paragraph WITH JUST THE IMAGE after the description,
# and add a caption below it.

# But extracting a drawing from a paragraph and putting it in a new one
# requires deepcopying the XML and modifying the paragraph. This is complex.

# Simpler approach: 
# For 5.2-5.4: The images are at the END of the description paragraphs.
# Add centered captions below each image paragraph.
# For 5.1: Add a caption below the image paragraph.

# Actually, in the description text, it already says "(1) 图 5.1..."
# That IS the caption. The issue is just that the image isn't visually separated.

# Let me take a much simpler approach:
# Just add a centered caption paragraph below each image paragraph.
# The caption will be a bold, centered "图 5.X xxx图" text.

caption_map = {
    35: '图 5.1 商品类别平均价格条形图',   # Image-only paragraph
}

# For figures 5.2-5.4 that are embedded, the caption text is already
# at the start of the paragraph. I should add the image in a separate
# paragraph below, with a caption below that.

# Let me insert new paragraphs after each mixed description paragraph

# Work in reverse order (bottom to top) to maintain indices
# The mixed paragraphs are at indices 36, 43, 44 (shifting as we insert)

# First, for each mixed para, create an image-only paragraph + caption
# and insert after the para

actions = []  # Will process bottom-to-top

# Figure 5.4 - para [44] (or whatever index after all processing)
# Find it by searching for the text
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '图 5.4' in text:
        # Check if it has an embedded drawing
        drawings = p._element.findall('.//'+qn('w:drawing'))
        if drawings:
            actions.append(('move_out', p, '图 5.4 商品评分分布箱线图', drawings[0]))
    elif '图 5.3' in text:
        drawings = p._element.findall('.//'+qn('w:drawing'))
        if drawings:
            actions.append(('move_out', p, '图 5.3 价格与评分关系散点图', drawings[0]))
    elif '图 5.2' in text:
        drawings = p._element.findall('.//'+qn('w:drawing'))
        if drawings:
            actions.append(('move_out', p, '图 5.2 主要商品类别分布占比饼图', drawings[0]))

print(f'\nActions to perform: {len(actions)}')
for a in actions:
    print(f'  {a[1].text[:40]}... -> para after with image + caption')

# Process from bottom to top
for action in reversed(actions):
    _, para, caption_text, drawing_elem = action
    
    # Deep copy the drawing element
    drawing_copy = deepcopy(drawing_elem)
    
    # Create image-only paragraph with centered alignment
    img_para = OxmlElement('w:p')
    img_pPr = OxmlElement('w:pPr')
    img_jc = OxmlElement('w:jc')
    img_jc.set(qn('w:val'), 'center')
    img_pPr.append(img_jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    img_pPr.append(spacing)
    img_para.append(img_pPr)
    img_para.append(drawing_copy)
    
    # Create caption paragraph
    cap_para = make_caption_para(caption_text)
    
    # Insert in order: caption after image after original paragraph
    para._element.addnext(cap_para)
    cap_para.addprevious(img_para)
    
    # Remove the drawing from the original paragraph
    # The drawing might be nested inside w:r (run), not directly under the para
    parent = drawing_elem.getparent()
    if parent is not None:
        parent.remove(drawing_elem)
        # If parent (w:r) is now empty, remove it too
        if len(parent) == 0:
            grandparent = parent.getparent()
            if grandparent is not None:
                grandparent.remove(parent)
    
    print(f'  Done: {caption_text}')

# Also add caption for Figure 5.1
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//'+qn('w:drawing'))
    if drawings and not p.text.strip():
        cap_para = make_caption_para('图 5.1 商品类别平均价格条形图')
        p._element.addnext(cap_para)
        print(f'  Added caption for 图 5.1 at para {i}')
        break

# ============================================
# STEP 3: Save
# ============================================
doc.save(r'C:\Users\KUHN\Desktop\1.docx')
print('\n=== SAVED ===')

# Final verification
doc2 = Document(r'C:\Users\KUHN\Desktop\1.docx')
print(f'\n=== FINAL STRUCTURE ===')
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()[:80] if p.text.strip() else ''
    has_img = '[IMG]' if len(p._element.findall('.//'+qn('w:drawing'))) > 0 else ''
    if p.style.name.startswith('Heading') or text or has_img:
        style_info = f'[{p.style.name}]'
        print(f'[{i:3d}] {style_info:20s} {text} {has_img}')
