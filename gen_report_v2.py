#!/usr/bin/env python3
"""
Generate experiment report with original cover page layout preserved.
Cover page: 软件学院 → 分布式系统原理与应用 → 课程作业 → info table
Body: content from 111.txt
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def make_p(text, size=12, bold=False, center=False, font_name='宋体', spacing_before=0, spacing_after=0):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if spacing_before:
        p.paragraph_format.space_before = Pt(spacing_before)
    if spacing_after:
        p.paragraph_format.space_after = Pt(spacing_after)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return p

# ═══════════════════════════════════════════════
# COVER PAGE - matching original format
# ═══════════════════════════════════════════════

# Spacer
for _ in range(3):
    doc.add_paragraph()

# 软件学院
make_p('软件学院', size=22, bold=True, center=True, font_name='黑体', spacing_before=12, spacing_after=6)

# Blank line
doc.add_paragraph()

# 分布式系统原理与应用
make_p('分布式系统原理与应用', size=20, bold=True, center=True, font_name='黑体', spacing_before=6, spacing_after=6)

# 课程作业
make_p('课程作业', size=18, bold=True, center=True, font_name='黑体', spacing_before=6, spacing_after=18)

# Info table (matching original layout: 作业题目 row, 姓名+学号+班级 row, 日期+导师签名 row)
info_table = doc.add_table(rows=3, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(4.5)

# Row 0: 作业题目 (spanning 3 cols)
# Merge cells 2,3 into cell 0 for a 2-col look on row 0
row0 = info_table.rows[0]
row0.cells[1].merge(row0.cells[2])
row0.cells[1].merge(row0.cells[3])

p00 = row0.cells[0].paragraphs[0]
p00.alignment = WD_ALIGN_PARAGRAPH.CENTER
r00 = p00.add_run('作业题目')
r00.bold = True
r00.font.size = Pt(12)

p01 = row0.cells[1].paragraphs[0]
p01.alignment = WD_ALIGN_PARAGRAPH.CENTER
r01 = p01.add_run('基于Redis分布式缓存的秒杀系统的实现')
r01.font.size = Pt(12)

# Row 1: 姓名 | 刘凯 | 学号 | 2310770132 | 班级 | 软件231
# We need 6 cells for this. Let me restructure: 
# Actually the original had: 姓名 刘凯    学号 2310770132 班级 软件231
# With 4 cols: [label1][value1][label2][value2]
# So row1 needs 6 cells for 姓名 刘凯 学号 2310770132 班级 软件231
# But we have 4 cols. Let me change approach.

# Let me redo: 3 rows, 6 cols
# Remove the table and recreate
info_table._element.getparent().remove(info_table._element)

info_table = doc.add_table(rows=3, cols=6)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

# Col widths: label=2cm, value=3cm, label=2cm, value=3cm, label=2cm, value=3cm
col_widths = [2.0, 3.5, 2.0, 3.5, 2.0, 3.5]
for row in info_table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = Cm(w)

def set_cell(table, row, col, text, bold=False, size=12, center=True):
    cell = table.cell(row, col)
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)

# Row 0: 作业题目 (merge cols 1-5)
info_table.cell(0, 0).merge(info_table.cell(0, 5))
# But that would make it one big cell. Let me instead merge 1-5.
info_table.cell(0, 0).merge(info_table.cell(0, 5))
# Now row 0 has one merged cell...

# Hmm, this is getting complicated. Let me use a simpler approach with 2 rows and then a merge.

# OK scrap the merged approach. Let me use a 4-column table for the info rows and handle the first row specially.

info_table._element.getparent().remove(info_table._element)

# New approach: single column table for first row, then 4-col for second row
# Actually the simplest approach matching the original layout:
# Use a single table with specific merges

info_table = doc.add_table(rows=4, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

# Row 0: 作业题目 (merged across 4 cols)
info_table.cell(0, 0).merge(info_table.cell(0, 3))
set_cell(info_table, 0, 0, '作业题目', bold=True, size=12, center=True)
set_cell(info_table, 0, 0, '：基于Redis分布式缓存的秒杀系统的实现', size=12, center=False)

# Row 1: 姓名 刘凯 | 学号 2310770132
info_table.cell(1, 0).merge(info_table.cell(1, 1))
info_table.cell(1, 2).merge(info_table.cell(1, 3))
set_cell(info_table, 1, 0, '姓名：刘凯', size=12, center=True)
set_cell(info_table, 1, 2, '学号：2310770132', size=12, center=True)

# Row 2: 班级 软件231
info_table.cell(2, 0).merge(info_table.cell(2, 1))
info_table.cell(2, 2).merge(info_table.cell(2, 3))
set_cell(info_table, 2, 0, '班级：软件231', size=12, center=True)

# Row 3: 日期 (blank) | 导师签名 (blank)
info_table.cell(3, 0).merge(info_table.cell(3, 1))
info_table.cell(3, 2).merge(info_table.cell(3, 3))
set_cell(info_table, 3, 0, '日期：          ', size=12, center=True)
set_cell(info_table, 3, 2, '导师签名：          ', size=12, center=True)

# Set column widths for the 4-col table
for row in info_table.rows:
    row.cells[0].width = Cm(6)
    row.cells[1].width = Cm(4)
    row.cells[2].width = Cm(4)
    row.cells[3].width = Cm(4)

# Add empty paragraphs after cover table
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Page break - separate cover from content
doc.add_page_break()

# ═══════════════════════════════════════════════
# BODY CONTENT from 111.txt
# ═══════════════════════════════════════════════

# Read 111.txt
txt_path = os.path.expandvars(r'%USERPROFILE%\Desktop\111.txt')
with open(txt_path, 'r', encoding='utf-8') as f:
    body_text = f.read()

# Split by sections and process
lines = body_text.split('\n')
current_section = None

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # Detect section headers
    if line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or line.startswith('四、') or line.startswith('五、'):
        # Major section header
        p = make_p(line, size=16, bold=True, font_name='黑体', spacing_before=18, spacing_after=10)
        continue
    
    # Detect sub-section headers (1.1, 2.1, etc.)
    if line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.'):
        if len(line) < 20:  # Short enough to be a sub-header
            p = make_p(line, size=14, bold=True, font_name='黑体', spacing_before=12, spacing_after=6)
            continue
    
    # Regular body text
    add_body(line)

add_body('（注：各成果展示部分请自行插入对应的功能截图。）')

# ── Save ──
output_path = os.path.expandvars(r'%USERPROFILE%\Desktop\分布式系统原理与应用-软件231-2310770132-刘凯.docx')
doc.save(output_path)
print(f'OK: report saved to {output_path}')
