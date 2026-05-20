from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document(r'C:\Users\KUHN\Desktop\1.docx')

# Check images
print('=== IMAGES ===')
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        print(f'  {rel.target_ref} -> {rel.reltype}')

# Check for inline images
image_count = 0
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//'+qn('w:drawing'))
    if drawings:
        for d in drawings:
            blip = d.find('.//'+qn('a:blip'))
            if blip is not None:
                embed = blip.get(qn('r:embed'))
                pref = p.text[:40] if p.text else '(empty)'
                print(f'  Para [{i}]: image embed={embed}, text="{pref}"')
                image_count += 1

print(f'\nTotal images found: {image_count}')

# Check paragraph count by section
print('\n=== DOCUMENT STRUCTURE ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()[:80] if p.text.strip() else ''
    if p.style.name.startswith('Heading') or text:
        style_info = f'[{p.style.name}]'
        has_img = ' [IMG]' if len(p._element.findall('.//'+qn('w:drawing'))) > 0 else ''
        print(f'[{i:3d}] {style_info:20s} {text}{has_img}')

# Check font of first few content paragraphs
print('\n=== FONT CHECK (sample) ===')
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Normal' and p.text.strip():
        for r in p.runs:
            font_cn = r._element.find(qn('w:rPr'))
            if font_cn is not None:
                rFonts = font_cn.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'), '?')
                    ascii_f = rFonts.get(qn('w:ascii'), '?')
                    sz = r.font.size
                    print(f'  [{i}] CN={ea} EN={ascii_f} size={sz} text="{r.text[:30]}"')
                    break
        if i > 35:  # Just first few
            break

# Check page margins
print('\n=== PAGE SETUP ===')
for si, section in enumerate(doc.sections):
    print(f'  Section {si}:')
    print(f'    Width: {section.page_width}, Height: {section.page_height}')
    print(f'    Margins: L={section.left_margin} R={section.right_margin} T={section.top_margin} B={section.bottom_margin}')
