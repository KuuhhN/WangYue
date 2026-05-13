#!/usr/bin/env python3
"""Generate the distributed experiment report from 111.txt content into proper .docx"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5

# Helper: add a heading-style paragraph
def add_heading_text(text, level=1, font_size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    return p

def add_sub_heading(text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level == 2 else Pt(12)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(4)
    return p

# ═══════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_heading_text('软件学院', level=0, font_size=22)
doc.add_paragraph()

add_heading_text('分布式系统原理与应用', level=0, font_size=20)
doc.add_paragraph()

add_heading_text('课程作业', level=0, font_size=18)
doc.add_paragraph()
doc.add_paragraph()

# Info table for cover
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

info_data = [
    ('作业题目', '基于Redis分布式缓存的秒杀系统的实现'),
    ('姓名', '刘凯'),
    ('学号', '2310770132'),
    ('班级', '软件231'),
    ('日期', '2026年5月'),
]

for i, (label, value) in enumerate(info_data):
    cell0 = info_table.cell(i, 0)
    cell1 = info_table.cell(i, 1)
    # Style label cell
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(label)
    r0.bold = True
    r0.font.size = Pt(14)
    # Style value cell
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(value)
    r1.font.size = Pt(14)

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(10)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 一、项目概况
# ═══════════════════════════════════════════════
add_heading_text('一、项目概况', level=1, font_size=16)

# 1.1 项目背景
add_sub_heading('1.1 项目背景')

add_body(
    '在电子商务与互联网技术飞速发展的时代背景下，电商平台的"双十一"、"618"等大规模营销活动已成为常态。'
    '在这些活动中，秒杀作为一种极具吸引力的营销手段，能够瞬间聚集海量用户流量。'
    '然而，秒杀活动具备典型的"瞬时高并发"特点，大量用户的读写请求在同一时间窗口内集中爆发。'
    '传统的单体架构因其资源限制，在高并发场景下会直接产生冲突：数据库（如MySQL）的连接池极易枯竭、'
    '中间件负载飙升，最终导致系统雪崩乃至完全宕机。'
)
add_body(
    '本项目的主要目标就是为了解决高并发场景下的系统性能瓶颈与数据一致性问题。'
    '通过设计并实现一个稳定可靠的分布式秒杀系统架构，一方面保障营销活动的公平公正性，'
    '提升用户的购物体验，避免因系统崩溃导致企业声誉受损和经济损失；'
    '另一方面，深入实践分布式缓存的核心特性——秒杀系统作为分布式架构的"试金石"，'
    '通过本项目全链路实战，能够深刻理解并掌握分布式系统架构中CAP理论的实际应用、'
    '数据的预加载策略以及原子性操作的实现等核心难点。'
    '这既是对理论知识的有效检验，也是从单体应用到分布式架构演进的必经之路。'
)

# 1.2 功能介绍
add_sub_heading('1.2 功能介绍')

add_body(
    '本项目围绕原生的电商秒杀活动的核心业务闭环，实现了从用户认证到商品展示、'
    '高并发秒杀、再到异步通知落地的全栈功能。系统采用前后端分离架构，主要模块划分如下：'
)

# Module table
mod_table = doc.add_table(rows=5, cols=2)
mod_table.style = 'Light Shading Accent 1'
mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
mod_headers = ['模块名称', '功能描述']
mod_data = [
    ('统一认证模块', '负责拦截非法访问与未登录的秒杀请求，用户需通过前台页面登录进行身份确认，响应底层用户信息校验。'),
    ('商品展示与秒杀模块', '动态渲染秒杀活动商品，系统预加载预热数据，前台通过模板引擎展示商品名称、图片占位图、原价以及极具吸引力的秒杀价（¥0.01元），同时页面动态展示Redis同步的实时剩余库存，并提供购买入口。'),
    ('分布式库存处理模块（核心）', '该模块是系统的减压重地。系统摒弃了传统数据库直接扣减库存的方式，转而利用Redis的单线程模型和高效内存处理优势，在缓存层面完成库存扣减。'),
    ('异步通知与日志记录模块', '用户在活动中秒杀成功后，系统生成异步任务（延迟执行），将成功订单写入日志并持久化到数据库，同时记录系统运行状态日志。'),
]

for i, (header, _) in enumerate([mod_headers]):
    for j, text in enumerate(mod_headers):
        cell = mod_table.cell(i, j)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)

for i, (mod_name, mod_desc) in enumerate(mod_data):
    cell0 = mod_table.cell(i+1, 0)
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(mod_name)
    r0.font.size = Pt(11)

    cell1 = mod_table.cell(i+1, 1)
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(mod_desc)
    r1.font.size = Pt(11)

mod_table.cell(0, 0).width = Cm(5)
mod_table.cell(0, 1).width = Cm(9)

doc.add_paragraph()  # spacing

add_body(
    '系统数据表结构设计共有8张核心表：goods（商品信息表）、order_info（订单流水表）、'
    'payment_record（支付状态记录）、seckill_goods（秒杀特定商品表）、seckill_order（秒杀成功用户记录表）、'
    'stock_log（库存变动明细）、sys_log（系统异常与运行日志表）以及 user（注册用户表）。'
)
add_body(
    '前端页面共5个核心页面：login.html（统一登录门户）、list.html（秒杀商品列表页）、'
    'detail.html（商品详情页）、result.html（用户秒杀结果展示页）、order.html（我的订单页）。'
    '前端采用Bootstrap + Thymeleaf模板渲染，确保页面美观且响应式良好。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 二、关键技术
# ═══════════════════════════════════════════════
add_heading_text('二、关键技术', level=1, font_size=16)

# 2.1 技术选型介绍
add_sub_heading('2.1 技术选型介绍')

add_body(
    '本项目技术栈采用当前主流的 Spring Boot 3.x 微服务框架进行开发，其核心优势在于通过'
    '自动化配置与约定大于配置的设计理念，开发者可快速构建并启动独立的生产级应用。'
    '持久层（DAO）选用 MyBatis-Plus，在原生 MyBatis 基础上简化了大量的样板代码，'
    '显著提升了对 MySQL 数据库的增删改查开发效率。'
)
add_body(
    '在分布式架构选型上，本项目核心的高并发流量压力由 Redis 分布式缓存中间件承载。'
    '实验环境方面，通过 VMware 虚拟机搭载 AlmaLinux 操作系统，并在其上部署 Redis 与 MySQL，'
    '模拟真实企业级前后端分离的系统架构。'
    '这一架构设计的核心指导思想是：尽力将全部读写操作压入内存层——因为磁盘 I/O 与 MySQL 的'
    '读写速度远远不及内存中 Redis 的处理速度——仅在库存逻辑校验通过后，'
    '才将最终持久化请求下发至数据库，以此达到保护底层系统的目的。'
)

# 2.2 某关键技术
add_sub_heading('2.2 关键技术：基于Redis的库存预加热与原子扣减')

add_body(
    '在本项目中，Redis分布式缓存技术贯穿于整个秒杀流程的核心环节，其技术实现主要体现于'
    '"库存预加热"和"原子化操作"两个维度。'
)

add_body(
    '首先是基于预加热技术的缓存预热。在传统架构设计中，用户首次访问秒杀页面时需要穿透式'
    '地从数据库获取数据，这在秒杀瞬时会导致"缓存击穿"问题。'
    '为此，本项目在系统启动阶段，利用Spring的 @PostConstruct 注解与初始化机制，'
    '通过依赖注入关系，从数据库中查询所有即将参与秒杀活动的商品信息，'
    '再通过系统内部通信将这些商品的库存数量转化为字符串格式，注入到远程Linux节点上的Redis内存中。'
    '通过预加热机制，秒杀时刻到来时的所有查询请求将直接命中Redis缓存层，'
    '响应时间从传统的毫秒级进一步压缩至微秒级，数据库连接的竞争压力也被从根本上释放。'
)

add_body(
    '其次是基于 decrement 原子扣减操作的精准限流。在库存处理这一业务公认的分布式瓶颈场景中，'
    '传统MySQL使用行锁解决并发问题，但其数据库连接池的有限性在超高并发下会成为新的瓶颈。'
    '本项目巧妙利用Redis底层采用单线程处理模型的天然原子性优势。'
    '当用户秒杀请求到达时，系统直接调用 stringRedisTemplate.opsForValue().decrement() 方法'
    '对对应商品的库存数量进行减一操作。由于Redis的单线程排队机制，每一次减一操作都是绝对的'
    '原子行为，不会被任何其他线程打断。系统仅判断扣减后的返回值——若返回值小于0，说明商品在此刻已被秒罄，'
    '系统立即向前端返回失败响应并执行回滚；只有返回值大于等于0的扣减请求，'
    '才被允许通过，继续进入数据库异步写入订单环节。'
    '这一设计在保证不超卖的前提下，实现了高性能的库存处理。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 三、实施步骤
# ═══════════════════════════════════════════════
add_heading_text('三、实施步骤', level=1, font_size=16)

add_body(
    '本系统的开发严格遵循工程化规范，从底层环境搭建到业务逻辑编写再到前端闭环，'
    '主要分为四个实施步骤：'
)

add_sub_heading('第一步：分布式缓存连接与参数配置')
add_body(
    '项目的基石是确保Spring Boot架构顺利连接到部署在远程Linux节点上的分布式缓存系统。'
    '在核心配置文件 application.yml 中进行严格的层级配置，除了常规数据库连接信息外，'
    '重点配置Redis连接参数，指定 spring.data.redis.host 为服务器固定IP地址，'
    '绑定默认的6379端口号，确保应用启动后能与Redis服务建立稳定的TCP连接。'
)

add_sub_heading('第二步：利用Bean生命周期的分布式预加热控制')
add_body(
    '连接通过后，开始编写秒杀核心业务类 SeckillService.java。利用 @PostConstruct 注解'
    '编写 initStock() 初始化方法，系统启动时自动全量加载MySQL数据库中的商品库存数据，'
    '通过 stringRedisTemplate.opsForValue().set() 方法，'
    '将商品ID拼接特定前缀后作为Redis键，对应的库存数量作为值，批量注入到Redis中。'
    '这一步骤确保了秒杀开始时所有库存数据已驻留在内存中。'
)

add_sub_heading('第三步：高并发秒杀接口的原子性扣减逻辑实现')
add_body(
    '在 SeckillService 中编写 executeSeckill(Long goodsId) 核心方法。'
    '接收到前端传入的商品ID后，直接使用 decrement 指令对Redis中对应的键值进行减一操作。'
    '判断返回值 stock：若 stock < 0，说明该商品已被秒罄，直接向前端返回失败消息；'
    '若 stock >= 0，则说明扣减成功，系统继续准备调用 updateById() '
    '真实修改MySQL数据库中的库存记录，完成最终数据持久化。'
)

add_sub_heading('第四步：MVC架构整合与前端业务闭环')
add_body(
    '在表示层，编写 GoodsController 作为前后端的交互映射。'
    '配置多个 @GetMapping 路径映射，包括登录成功后自动重定向、'
    '获取商品列表数据的接口 /list 以及执行秒杀的核心接口 /buy。'
    '同时，在 templates 目录下编写5个HTML文件（登录页、列表页、详情页、结果页、订单页），'
    '使用Bootstrap + Thymeleaf模板引擎渲染动态数据，完成完整的电商秒杀业务流程闭环。'
    '最终实现从前端页面发起请求，到后端Redis缓存层承接、'
    '原子扣减、限流、再异步写入数据库的完整数据流链路。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 四、成果展示
# ═══════════════════════════════════════════════
add_heading_text('四、成果展示', level=1, font_size=16)

add_body(
    '系统运行后，各功能均达到预定设计目标，主要成果展示如下：'
)

add_sub_heading('1. 系统统一登录页展示')
add_body(
    '用户访问系统根目录时，会被拦截并重定向到登录页面，确保系统的安全性。'
    '登录页采用Bootstrap框架设计，界面简洁美观，包含用户账号密码输入框及登录按钮。'
    '（此处插入 login.html 登录页面的截图）'
)

add_sub_heading('2. 秒杀商品列表与分布式预热展示')
add_body(
    '系统启动时，后台通过分布式通信成功将MySQL数据库数据精准转存至Redis内存中。'
    '前台列表页成功获取所有参与秒杀的商品，展示商品名称、占位图及秒杀价¥0.01元。'
    '（此处插入 list.html 秒杀商品列表页面的截图）'
)

add_sub_heading('3. 实时库存同步与秒杀入口展示')
add_body(
    '进入单品详情页后，系统能够准确展示商品的原始价格、秒杀价格以及通过前端轮询从Redis'
    '拉取的实时剩余库存数量，为用户决策提供透明的库存信息。'
    '（此处插入 detail.html 商品详情页面的截图）'
)

add_sub_heading('4. 高并发秒杀与结果展示')
add_body(
    '当用户在秒杀倒计时结束后点击购买时，页面瞬间跳转至结果展示页。'
    '后台成功运用分布式缓存实现了库存扣减与限流，有效保护了底层数据库。'
    '（此处插入 result.html 秒杀成功/失败结果页面的截图）'
)

add_sub_heading('5. 个人订单页展示')
add_body(
    '通过Redis内存级扣减和限流，成功在底层数据库中生成正确的订单记录，完成业务闭环。'
    '订单页完整展示用户已下单商品的详细信息。'
    '（此处插入 order.html 我的订单页面的截图）'
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 五、文档感悟
# ═══════════════════════════════════════════════
add_heading_text('五、文档感悟', level=1, font_size=16)

# 5.1 技术掌握情况
add_sub_heading('5.1 技术掌握情况')

add_body(
    '通过本次课程项目的实战开发，我对目前业界主流的 Spring Boot 3.x 技术体系有了更为深刻的理解和运用，'
    '熟练掌握了MVC架构中各层级的职责划分与流转逻辑。在分布式缓存方面，'
    '成功掌握了如何在Linux虚拟机环境中部署并配置Redis中间件，'
    '特别是对原子性操作（decrement）的底层实现与实战应用有了深刻认识。'
)
add_body(
    '然而，项目实现中也暴露出了一些不足之处：当前系统虽然实现了核心的分布式缓存功能，'
    '但异步通知与日志持久化仍然采用的是同步处理思路，未引入真正的消息队列（如 RabbitMQ 或 Kafka）'
    '来承载削峰填谷的职责；此外，在分布式事务一致性保障与系统监控可观测性方面尚属空白，'
    '这将是后续学习中的重点攻关方向。'
)

# 5.2 课程技术收获
add_sub_heading('5.2 课程技术收获')

add_body(
    '在《分布式系统原理与应用》这门课程中，最大的收获是从传统的"单体思维"向"集群架构思维"的跨越。'
    '理论课上学习的 CAP 理论等抽象概念，通过项目实训得以具象化——在秒杀的高并发场景下，'
    '既需要保证数据的一致性（C），也要通过 Redis 确保系统的高可用性（A）与响应速度，'
    '同时理解了分布式系统中"空间换时间"的精髓思想——将昂贵的磁盘 I/O 转化为廉价的内存操作。'
    '从理论到实践，对分布式系统中的数据分片、负载均衡、故障转移等核心概念都有了更立体的认知。'
)

# 5.3 学习过程感悟
add_sub_heading('5.3 学习过程感悟')

add_body(
    '在具体的编码实践和非技术层面，最深切的体会是"细节决定成败"。在开发初期，'
    '由于 application.yml 文件中一个参数配置错误（缩进格式），导致应用启动后Redis连接失败并报错；'
    '也因为在编写库存扣减逻辑时忽略了递减操作的边界判断，导致了测试环境下的超卖问题。'
    '这些坑经验进一步磨练了我的问题定位能力和代码严谨性。'
    '通过逐行分析错误日志，我学会了不盲目搜索、而是能够顺藤摸瓜地准确定位代码根源。'
    '这次全链路实战的"从零到一"经历，不仅提升了我的编码功底，'
    '更培养了我面对复杂系统问题时保持冷静、系统性思考和解决 Bug 的工程师心态。'
)

doc.add_paragraph()
add_body_no_indent('（注：以上红色文字要求在完成后统一删除。如有配图需求，请自行插入对应的功能截图。）')

# ── Save ──
output_path = os.path.expandvars(r'%USERPROFILE%\Desktop\分布式系统原理与应用-软件231-2310770132-刘凯.docx')
doc.save(output_path)
print(f'OK: report saved to {output_path}')
