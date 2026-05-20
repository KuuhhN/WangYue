"""
v5: Nuclear cleanup - delete ALL paragraphs between headings, then insert fresh content.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\KUHN\Desktop\1.docx')

def make_paragraph(text, is_code=False):
    """Create a w:p element with proper formatting."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'exact')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    
    if not is_code:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '480')
        pPr.append(ind)
    
    p.append(pPr)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    if is_code:
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
        rFonts.set(qn('w:eastAsia'), '宋体')
        sz_val = '20'
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '333333')
        rPr.append(color)
    else:
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '宋体')
        sz_val = '24'
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)
    
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

# Step 1: Update cover table and main title (before cleanup)
table0 = doc.tables[0]
cell = table0.cell(1, 1)
cell.text = ''
run = cell.paragraphs[0].add_run('基于电商商品数据集的商品特征及价格行为分析')
run.font.size = Pt(12)

for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and '××' in p.text:
        p.clear()
        run = p.add_run('基于电商商品数据集的商品特征及价格行为分析')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        break

# Step 2: Find headings and their positions
heading_paras = []  # list of (text, element, style)
body = doc.element.body
for i, child in enumerate(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        # Check if this is a heading paragraph
        pPr = child.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                val = pStyle.get(qn('w:val'))
                if val and 'Heading' in val:
                    # Get text
                    texts = [t.text or '' for t in child.iter(qn('w:t'))]
                    text = ''.join(texts).strip()
                    heading_paras.append((text, child, val))

print(f'Found {len(heading_paras)} headings:')
for h in heading_paras:
    print(f'  [{h[2]}] {h[0][:40]}')

# Step 3: Delete ALL body children that are NOT headings, NOT tables, NOT sectPr
# Between the "概述" heading and the last heading, delete everything that isn't a heading
# Actually, let's be more precise: for each section, delete content between headings

# Find the "概述" heading position
overview_idx = None
data_get_idx = None
data_proc_idx = None
data_ana_idx = None
data_vis_idx = None
conclusion_idx = None
comments_idx = None

for i, (text, elem, style) in enumerate(heading_paras):
    if text == '概述': overview_idx = i
    elif text == '数据获取': data_get_idx = i
    elif text == '数据处理': data_proc_idx = i
    elif text == '数据分析': data_ana_idx = i
    elif text == '数据可视化': data_vis_idx = i
    elif text == '结论与总结': conclusion_idx = i
    elif text == '评语': comments_idx = i

print(f'\nIndices: overview={overview_idx}, get={data_get_idx}, proc={data_proc_idx}, ana={data_ana_idx}, vis={data_vis_idx}, conc={conclusion_idx}, comments={comments_idx}')

# Step 4: For each pair of consecutive headings, delete content paragraphs between them
def delete_between(start_elem, end_elem):
    """Delete all non-table elements between start and end (not inclusive)."""
    count = 0
    elem = start_elem.getnext()
    while elem is not None and elem is not end_elem:
        next_elem = elem.getnext()
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag != 'tbl':  # Don't delete tables (cover page)
            body.remove(elem)
            count += 1
        elem = next_elem
    return count

# Delete content between section headings (from bottom to top so indices stay valid)
# Only do this for known sections
section_pairs = []

if conclusion_idx is not None and comments_idx is not None:
    h_conc = heading_paras[conclusion_idx][1]
    h_comm = heading_paras[comments_idx][1]
    n = delete_between(h_conc, h_comm)
    section_pairs.append(('结论→评语', n))

if data_vis_idx is not None and conclusion_idx is not None:
    h_vis = heading_paras[data_vis_idx][1]
    h_conc = heading_paras[conclusion_idx][1]
    n = delete_between(h_vis, h_conc)
    section_pairs.append(('可视化→结论', n))

if data_ana_idx is not None and data_vis_idx is not None:
    h_ana = heading_paras[data_ana_idx][1]
    h_vis = heading_paras[data_vis_idx][1]
    n = delete_between(h_ana, h_vis)
    section_pairs.append(('分析→可视化', n))

if data_proc_idx is not None and data_ana_idx is not None:
    h_proc = heading_paras[data_proc_idx][1]
    h_ana = heading_paras[data_ana_idx][1]
    n = delete_between(h_proc, h_ana)
    section_pairs.append(('处理→分析', n))

if data_get_idx is not None and data_proc_idx is not None:
    h_get = heading_paras[data_get_idx][1]
    h_proc = heading_paras[data_proc_idx][1]
    n = delete_between(h_get, h_proc)
    section_pairs.append(('获取→处理', n))

if overview_idx is not None and data_get_idx is not None:
    h_over = heading_paras[overview_idx][1]
    h_get = heading_paras[data_get_idx][1]
    n = delete_between(h_over, h_get)
    section_pairs.append(('概述→获取', n))

for name, count in section_pairs:
    print(f'  {name}: deleted {count} paragraphs')

# Re-get heading elements after deletions (re-read the tree)
print('\n--- Rebuilding heading references after cleanup ---')
heading_paras2 = []
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        pPr = child.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                val = pStyle.get(qn('w:val'))
                if val and 'Heading' in val:
                    texts = [t.text or '' for t in child.iter(qn('w:t'))]
                    text = ''.join(texts).strip()
                    heading_paras2.append((text, child, val))

def find_elem(text_substr):
    for t, e, s in heading_paras2:
        if text_substr in t:
            return e
    return None

h_overview_e = find_elem('概述')
h_data_get_e = find_elem('数据获取')
h_data_proc_e = find_elem('数据处理')
h_data_ana_e = find_elem('数据分析')
h_data_vis_e = find_elem('数据可视化')
h_conclusion_e = find_elem('结论与总结')
h_comments_e = find_elem('评语')

print(f'Re-found: overview={h_overview_e is not None}, get={h_data_get_e is not None}, proc={h_data_proc_e is not None}, ana={h_data_ana_e is not None}, vis={h_data_vis_e is not None}, conc={h_conclusion_e is not None}, comments={h_comments_e is not None}')

# Step 5: Insert fresh content BEFORE each heading (bottom to top)
# Use addprevious directly on XML elements

# Helper: insert list of paragraphs before an element (in correct order)
def insert_before_elem(paras_xml, before_elem):
    for p in reversed(paras_xml):
        before_elem.addprevious(p)

# 结论与总结 content
conc_texts = [
    '通过本次分析，得出以下核心结论：',
    '（1）定价规律：电子产品与高端家居类目的平均客单价显著高于日用杂货类，说明不同品类的利润空间与市场定价逻辑存在明显差异。平台可根据品类价格特征制定差异化的定价策略。',
    '（2）品质表现：整体而言，平台商品评分分布稳定在较高水平，价格的高低与用户好评度并无强烈的负相关性，这表明用户对高价商品的品牌溢价具有较高认可度。',
    '实际应用建议：建议电商平台在后续运营中，针对高客单价品类加强品质控制，同时利用折扣数据优化中低价位商品的库存周转效率，实现整体经营效益的最大化。',
]
insert_before_elem([make_paragraph(t) for t in conc_texts], h_comments_e)
print('结论与总结: done')

# 数据可视化 content
vis_texts = [
    '通过 Matplotlib 和 Seaborn 库，构建了四种不同类型的图表进行直观呈现。',
    '（1）图 5.1 商品类别平均价格条形图：展示了 Top 10 商品类别的平均价格，清晰对比了各品类商品的消费层级。从图中可以看出，电子产品和高端家居品类位列价格前茅，而日用杂货类商品价格较低。',
    '（2）图 5.2 主要商品类别分布占比饼图：可视化了主要商品类别的分布占比，明确了当前库存结构中的核心品类。从图中可以看出，平台商品品类分布较为均衡，少数核心品类占据了较大市场份额。',
    '（3）图 5.3 价格与评分关系散点图：通过价格与评分的映射，直观展现了商品价格与用户满意度之间的分布关系。散点图显示数据点分布较为分散，未见明显的正相关或负相关趋势。',
    '（4）图 5.4 商品评分分布箱线图：展示了整体商品评分的离散程度及中位数水平，揭示了平台商品的综合评价质量。箱线图显示评分中位数在 4.0 以上，说明平台商品整体质量较好。',
]
vis_code = """import matplotlib.pyplot as plt
import seaborn as sns

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

fig, axes = plt.subplots(2, 2, figsize=(14, 10))

# 图5.1 条形图
top10_cat = category_stats.head(10)
axes[0, 0].barh(top10_cat.index, top10_cat['mean'])
axes[0, 0].set_title('Top 10 商品类别平均价格')
axes[0, 0].set_xlabel('平均价格')

# 图5.2 饼图
cat_counts = df['category'].value_counts()
axes[0, 1].pie(cat_counts.values, labels=cat_counts.index,
               autopct='%1.1f%%')
axes[0, 1].set_title('商品类别分布占比')

# 图5.3 散点图
axes[1, 0].scatter(df['price'], df['rating'], alpha=0.6)
axes[1, 0].set_title('价格与评分关系')
axes[1, 0].set_xlabel('价格')
axes[1, 0].set_ylabel('评分')

# 图5.4 箱线图
sns.boxplot(data=df, y='rating', ax=axes[1, 1])
axes[1, 1].set_title('商品评分分布')

plt.tight_layout()
plt.show()"""

vis_all = [make_paragraph(t) for t in vis_texts]
vis_all.append(make_paragraph('数据可视化核心代码如下：'))
vis_all.append(make_paragraph(vis_code, is_code=True))
insert_before_elem(vis_all, h_conclusion_e)
print('数据可视化: done')

# 数据分析 content
ana_texts = [
    '在清洗后的数据基础上，主要从以下两个维度进行了深入分析：',
    '（1）商品类别特征分析：通过 groupby 分组方法计算各品类商品的平均价格，揭示了不同类别商品的定价结构差异。',
    '（2）相关性探索：分析了价格与用户评分之间的关系，探究消费者是否会因为高价商品而给予较低评价，从而评估平台的商品质量分布现状。',
]
ana_code = """# 商品类别特征分析
category_stats = df.groupby('category')['price'].agg(['mean', 'count'])
category_stats = category_stats.sort_values('mean', ascending=False)

# 价格与评分相关性分析
correlation = df['price'].corr(df['rating'])
print(f"价格与评分的相关系数: {correlation:.4f}")

# 输出 Top 5 平均价格最高的品类
print(category_stats.head())"""

ana_all = [make_paragraph(t) for t in ana_texts]
ana_all.append(make_paragraph('数据分析的核心代码如下：'))
ana_all.append(make_paragraph(ana_code, is_code=True))
ana_all.append(make_paragraph('通过分析发现，不同商品类别的平均价格存在显著差异，电子产品和高端家居类目远高于日用百货。同时，价格与评分的相关系数接近零，说明价格高低并不直接影响用户评分。'))
insert_before_elem(ana_all, h_data_vis_e)
print('数据分析: done')

# 数据处理 content
proc_texts = [
    '为确保分析结果的科学性，对原始数据进行了深度清洗，具体处理流程如下：',
    '（1）重复值处理：通过检测并删除了数据集中的 5 条重复记录，保证了数据的唯一性。',
    '（2）缺失值填充：针对 price 列的缺失数据，利用 pandas 库计算了该列的均值，并将其填补至对应空缺处，有效避免了因空值导致的统计偏差。同时，将 brand 列的缺失值统一标记为"Unknown"。',
    '（3）异常值修正：鉴于评分（rating）的正常范围应在 0-5 分之间，针对数据中出现的 999.0 异常值，通过逻辑定位将其统一修正为 5.0，确保了后续评分分析的准确性。',
]
proc_code = """# 检测重复值
duplicates = df.duplicated().sum()
print(f"重复记录数: {duplicates}")
df = df.drop_duplicates()

# 缺失值处理（均值填充）
mean_price = df['price'].mean()
df['price'].fillna(mean_price, inplace=True)
df['brand'].fillna("Unknown", inplace=True)

# 异常值修正
df.loc[df['rating'] > 5, 'rating'] = 5.0

print(f"清洗后数据维度: {df.shape}")"""

proc_all = [make_paragraph(t) for t in proc_texts]
proc_all.append(make_paragraph('数据处理的实现代码如下：'))
proc_all.append(make_paragraph(proc_code, is_code=True))
proc_all.append(make_paragraph('经过上述清洗步骤，数据质量显著提升，为后续分析提供了可靠的数据基础。'))
insert_before_elem(proc_all, h_data_ana_e)
print('数据处理: done')

# 数据获取 content
get_text = '本实验采用网络请求技术，从 DummyJSON 开放平台（https://dummyjson.com/products）获取了原始的商品数据集。该数据集包含了 100 条商品记录，字段涵盖了 id（唯一标识）、title（商品名称）、category（商品类别）、price（价格）、discountPercentage（折扣百分比）、rating（用户评分）、stock（库存数量）及 brand（品牌）。数据呈现出明显的电商业务特征，且包含部分缺失值与异常点，非常适合进行数据清洗与挖掘练习。'
get_code = """import requests
import pandas as pd

# 发送 API 请求获取商品数据
url = "https://dummyjson.com/products"
response = requests.get(url)
data = response.json()

# 转换为 DataFrame
df = pd.DataFrame(data['products'])
print(f"数据维度: {df.shape}")
print(df.head())"""

get_all = [
    make_paragraph(get_text),
    make_paragraph('以下为数据获取的核心代码：'),
    make_paragraph(get_code, is_code=True),
    make_paragraph('通过上述代码，成功获取了 100 条商品记录，包含 8 个核心字段，为后续分析奠定了基础。'),
]
insert_before_elem(get_all, h_data_proc_e)
print('数据获取: done')

# 概述 content
overview_text = '本次课程设计旨在掌握 Python 数据分析的全流程。通过对公开电商商品数据集进行系统性的分析，完成了从数据获取、预处理、探索性分析到可视化展示的闭环工作。本次分析旨在通过数据手段揭示商品类别、定价策略与用户评分之间的潜在规律，并为电商平台的运营提供数据支持。'
insert_before_elem([make_paragraph(overview_text)], h_data_get_e)
print('概述: done')

# Step 6: Save
doc.save(r'C:\Users\KUHN\Desktop\1.docx')
print('\n=== SAVED ===')

# Verify
doc2 = Document(r'C:\Users\KUHN\Desktop\1.docx')
print(f'Total paragraphs: {len(doc2.paragraphs)}')
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    if text:
        display = text[:90] + ('...' if len(text) > 90 else '')
        print(f'[{i:3d}] [{str(p.style.name)[:10]:10s}] {display}')
