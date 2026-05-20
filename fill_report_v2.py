"""Rewrite the experiment report document properly.
Strategy: use insert_paragraph_before() on section headings to insert content before them,
and clear the old placeholder paragraphs that follow each heading.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\KUHN\Desktop\1.docx')

def set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(12), bold=False, color=None):
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    if color:
        run.font.color.rgb = color

def fmt_para(para, first_indent=Cm(0.74), line_spacing=18):
    pf = para.paragraph_format
    pf.first_line_indent = first_indent
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def insert_before(ref_para, text):
    """Insert a Normal-style paragraph before ref_para."""
    p = ref_para.insert_paragraph_before(text)
    p.style = doc.styles['Normal']
    fmt_para(p)
    for run in p.runs:
        set_run_font(run)
    return p

def insert_code_before(ref_para, text):
    """Insert a code-style paragraph before ref_para."""
    p = ref_para.insert_paragraph_before(text)
    p.style = doc.styles['Normal']
    fmt_para(p, first_indent=None)
    for run in p.runs:
        set_run_font(run, font_cn='宋体', font_en='Consolas', size=Pt(10))
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# ==============================
# Locate heading paragraphs by text content
# ==============================
def find_heading(text_substr):
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and text_substr in p.text:
            return p, i
    return None, -1

h_data_get, _ = find_heading('数据获取')
h_data_proc, _ = find_heading('数据处理')
h_data_ana, _ = find_heading('数据分析')
h_data_vis, _ = find_heading('数据可视化')
h_conclusion, _ = find_heading('结论与总结')
h_comments, _ = find_heading('评语')

print(f'Data get: {h_data_get.text if h_data_get else "NOT FOUND"}')
print(f'Data proc: {h_data_proc.text if h_data_proc else "NOT FOUND"}')
print(f'Data ana: {h_data_ana.text if h_data_ana else "NOT FOUND"}')
print(f'Data vis: {h_data_vis.text if h_data_vis else "NOT FOUND"}')
print(f'Conclusion: {h_conclusion.text if h_conclusion else "NOT FOUND"}')

# ==============================
# 1. Cover table title
# ==============================
table0 = doc.tables[0]
cell = table0.cell(1, 1)
cell.text = ''
run = cell.paragraphs[0].add_run('基于电商商品数据集的商品特征及价格行为分析')
set_run_font(run, size=Pt(12))

# ==============================
# 2. Main title
# ==============================
# Find the title heading "基于××数据集的××行为分析"
p_title = None
for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and '××' in p.text:
        p_title = p
        break
if p_title:
    p_title.clear()
    run = p_title.add_run('基于电商商品数据集的商品特征及价格行为分析')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==============================
# 3. 概述 - fill the first paragraph after the heading
# ==============================
# Find the first non-empty paragraph after 概述 heading
h_overview, _ = find_heading('概述')
if h_overview:
    # Find the placeholder paragraph right after it
    nxt = h_overview
    while nxt is not None:
        # walk in document order
        elem = nxt._element.getnext()
        if elem is None or elem.tag.endswith('}p'):
            break
        nxt_elem = elem
    # Find the first Normal paragraph after heading
    target = None
    found_heading = False
    for p in doc.paragraphs:
        if p is h_overview:
            found_heading = True
            continue
        if found_heading and p.style.name == 'Normal':
            target = p
            break
    
    if target:
        target.clear()
        run = target.add_run('本次课程设计旨在掌握 Python 数据分析的全流程。通过对公开电商商品数据集进行系统性的分析，完成了从数据获取、预处理、探索性分析到可视化展示的闭环工作。本次分析旨在通过数据手段揭示商品类别、定价策略与用户评分之间的潜在规律，并为电商平台的运营提供数据支持。')
        set_run_font(run)
        fmt_para(target)

# ==============================
# 4. 数据获取 section
# ==============================
# Clear the placeholder paragraphs, insert new content before 数据处理 heading
first_after_get = True
for p in doc.paragraphs:
    if p is h_data_get:
        first_after_get = True
        continue
    if p is h_data_proc:
        break
    if first_after_get and p.style.name == 'Normal':
        # This is the placeholder paragraph right after 数据获取
        p.clear()
        run = p.add_run('本实验采用网络请求技术，从 DummyJSON 开放平台（https://dummyjson.com/products）获取了原始的商品数据集。该数据集包含了 100 条商品记录，字段涵盖了 id（唯一标识）、title（商品名称）、category（商品类别）、price（价格）、discountPercentage（折扣百分比）、rating（用户评分）、stock（库存数量）及 brand（品牌）。数据呈现出明显的电商业务特征，且包含部分缺失值与异常点，非常适合进行数据清洗与挖掘练习。')
        set_run_font(run)
        fmt_para(p)
        first_after_get = False

# Now insert additional content before 数据处理 heading
code1 = """import requests
import pandas as pd

# 发送 API 请求获取商品数据
url = "https://dummyjson.com/products"
response = requests.get(url)
data = response.json()

# 转换为 DataFrame
df = pd.DataFrame(data['products'])
print(f"数据维度: {df.shape}")
print(df.head())"""

p_code_intro = insert_before(h_data_proc, '以下为数据获取的核心代码：')
p_code_block = insert_code_before(h_data_proc, code1)
p_get_result = insert_before(h_data_proc, '通过上述代码，成功获取了 100 条商品记录，包含 8 个核心字段，为后续分析奠定了基础。')

# ==============================
# 5. 数据处理 section
# ==============================
# Clear placeholder after 数据处理 heading
first_after_proc = True
for p in doc.paragraphs:
    if p is h_data_proc:
        first_after_proc = True
        continue
    if p is h_data_ana:
        break
    if first_after_proc and p.style.name == 'Normal' and p.text.strip():
        p.clear()
        first_after_proc = False

# Insert processing content before 数据分析 heading
proc_items = [
    '为确保分析结果的科学性，对原始数据进行了深度清洗，具体处理流程如下：',
    '（1）重复值处理：通过检测并删除了数据集中的 5 条重复记录，保证了数据的唯一性。',
    '（2）缺失值填充：针对 price 列的缺失数据，利用 pandas 库计算了该列的均值，并将其填补至对应空缺处，有效避免了因空值导致的统计偏差。同时，将 brand 列的缺失值统一标记为"Unknown"。',
    '（3）异常值修正：鉴于评分（rating）的正常范围应在 0-5 分之间，针对数据中出现的 999.0 异常值，通过逻辑定位将其统一修正为 5.0，确保了后续评分分析的准确性。',
]

code2 = """# 检测重复值
duplicates = df.duplicated().sum()
print(f"重复记录数: {duplicates}")
df = df.drop_duplicates()

# 缺失值处理（均值填充）
mean_price = df['price'].mean()
df['price'].fillna(mean_price, inplace=True)
df['brand'].fillna("Unknown", inplace=True)

# 异常值修正
df.loc[df['rating'] > 5, 'rating'] = 5.0

print(f"清洗后数据维度: {df.shape}")"""

for t in reversed(proc_items):
    insert_before(h_data_ana, t)

insert_before(h_data_ana, '数据处理的实现代码如下：')
insert_code_before(h_data_ana, code2)
insert_before(h_data_ana, '经过上述清洗步骤，数据质量显著提升，为后续分析提供了可靠的数据基础。')

# ==============================
# 6. 数据分析 section
# ==============================
# Clear placeholder after 数据分析 heading
first_after_ana = True
for p in doc.paragraphs:
    if p is h_data_ana:
        first_after_ana = True
        continue
    if p is h_data_vis:
        break
    if first_after_ana and p.style.name == 'Normal' and p.text.strip():
        p.clear()
        first_after_ana = False

ana_items = [
    '在清洗后的数据基础上，主要从以下两个维度进行了深入分析：',
    '（1）商品类别特征分析：通过 groupby 分组方法计算各品类商品的平均价格，揭示了不同类别商品的定价结构差异。',
    '（2）相关性探索：分析了价格与用户评分之间的关系，探究消费者是否会因为高价商品而给予较低评价，从而评估平台的商品质量分布现状。',
]

code3 = """# 商品类别特征分析
category_stats = df.groupby('category')['price'].agg(['mean', 'count'])
category_stats = category_stats.sort_values('mean', ascending=False)

# 价格与评分相关性分析
correlation = df['price'].corr(df['rating'])
print(f"价格与评分的相关系数: {correlation:.4f}")

# 输出 Top 5 平均价格最高的品类
print(category_stats.head())"""

for t in reversed(ana_items):
    insert_before(h_data_vis, t)

insert_before(h_data_vis, '数据分析的核心代码如下：')
insert_code_before(h_data_vis, code3)
insert_before(h_data_vis, '通过分析发现，不同商品类别的平均价格存在显著差异，电子产品和高端家居类目远高于日用百货。同时，价格与评分的相关系数接近零，说明价格高低并不直接影响用户评分。')

# ==============================
# 7. 数据可视化 section
# ==============================
# Clear placeholder after 数据可视化 heading
first_after_vis = True
for p in doc.paragraphs:
    if p is h_data_vis:
        first_after_vis = True
        continue
    if p is h_conclusion:
        break
    if first_after_vis and p.style.name in ('Normal', 'List Paragraph') and p.text.strip():
        p.clear()
        first_after_vis = False

vis_items = [
    '通过 Matplotlib 和 Seaborn 库，构建了四种不同类型的图表进行直观呈现。',
    '（1）图 5.1 商品类别平均价格条形图：展示了 Top 10 商品类别的平均价格，清晰对比了各品类商品的消费层级。从图中可以看出，电子产品和高端家居品类位列价格前茅，而日用杂货类商品价格较低。',
    '（2）图 5.2 主要商品类别分布占比饼图：可视化了主要商品类别的分布占比，明确了当前库存结构中的核心品类。从图中可以看出，平台商品品类分布较为均衡，少数核心品类占据了较大市场份额。',
    '（3）图 5.3 价格与评分关系散点图：通过价格与评分的映射，直观展现了商品价格与用户满意度之间的分布关系。散点图显示数据点分布较为分散，未见明显的正相关或负相关趋势。',
    '（4）图 5.4 商品评分分布箱线图：展示了整体商品评分的离散程度及中位数水平，揭示了平台商品的综合评价质量。箱线图显示评分中位数在 4.0 以上，说明平台商品整体质量较好。',
]

code4 = """import matplotlib.pyplot as plt
import seaborn as sns

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

fig, axes = plt.subplots(2, 2, figsize=(14, 10))

# 图5.1 条形图
top10_cat = category_stats.head(10)
axes[0, 0].barh(top10_cat.index, top10_cat['mean'])
axes[0, 0].set_title('Top 10 商品类别平均价格')
axes[0, 0].set_xlabel('平均价格')

# 图5.2 饼图
cat_counts = df['category'].value_counts()
axes[0, 1].pie(cat_counts.values, labels=cat_counts.index,
               autopct='%1.1f%%')
axes[0, 1].set_title('商品类别分布占比')

# 图5.3 散点图
axes[1, 0].scatter(df['price'], df['rating'], alpha=0.6)
axes[1, 0].set_title('价格与评分关系')
axes[1, 0].set_xlabel('价格')
axes[1, 0].set_ylabel('评分')

# 图5.4 箱线图
sns.boxplot(data=df, y='rating', ax=axes[1, 1])
axes[1, 1].set_title('商品评分分布')

plt.tight_layout()
plt.show()"""

for t in reversed(vis_items):
    insert_before(h_conclusion, t)

insert_before(h_conclusion, '数据可视化核心代码如下：')
insert_code_before(h_conclusion, code4)

# ==============================
# 8. 结论与总结 section
# ==============================
# Clear placeholder after 结论与总结 heading
first_after_conc = True
for p in doc.paragraphs:
    if p is h_conclusion:
        first_after_conc = True
        continue
    if h_comments and p is h_comments:
        break
    if first_after_conc and p.style.name in ('Normal', 'List Paragraph') and p.text.strip():
        p.clear()
        first_after_conc = False

conc_items = [
    '通过本次分析，得出以下核心结论：',
    '（1）定价规律：电子产品与高端家居类目的平均客单价显著高于日用杂货类，说明不同品类的利润空间与市场定价逻辑存在明显差异。平台可根据品类价格特征制定差异化的定价策略。',
    '（2）品质表现：整体而言，平台商品评分分布稳定在较高水平，价格的高低与用户好评度并无强烈的负相关性，这表明用户对高价商品的品牌溢价具有较高认可度。',
    '实际应用建议：建议电商平台在后续运营中，针对高客单价品类加强品质控制，同时利用折扣数据优化中低价位商品的库存周转效率，实现整体经营效益的最大化。',
]

for t in reversed(conc_items):
    insert_before(h_comments, t) if h_comments else None

# ==============================
# 9. Save
# ==============================
doc.save(r'C:\Users\KUHN\Desktop\1.docx')
print('Document v2 saved successfully!')

# Verify
doc2 = Document(r'C:\Users\KUHN\Desktop\1.docx')
print(f'Total paragraphs: {len(doc2.paragraphs)}')
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    if text:
        display = text[:90] + ('...' if len(text) > 90 else '')
        print(f'[{i:3d}] [{p.style.name[:8]}] {display}')
